"""Tests endpoint /documents/preview-change."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

import app.auth as authmod
import app.main as mainmod
from app.documents.preview_change import build_line_diff_html


@pytest.fixture(autouse=True)
def _loopback(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(authmod, "is_loopback_host", lambda host: True)


def _headers() -> dict[str, str]:
    return {"X-Internal-Secret": "desktop-dev-secret"}


def test_build_line_diff_html_marks_add_and_del() -> None:
    html = build_line_diff_html("ligne A\n", "ligne B\n")
    assert "wp-diff-add" in html
    assert "wp-diff-del" in html
    assert "ligne B" in html
    assert "ligne A" in html


def test_preview_change_text_file(tmp_path: Path) -> None:
    project = tmp_path / "project"
    project.mkdir()
    ws_data = tmp_path / "ws_data"
    ws_data.mkdir()
    (project / "note.md").write_text("ancien\nligne 2", encoding="utf-8")

    with TestClient(mainmod.app) as client:
        resp = client.post(
            "/documents/preview-change",
            json={
                "workspace_data_dir": str(ws_data),
                "project_path": str(project),
                "file_path": "note.md",
                "proposed_content": "nouveau\nligne 2",
            },
            headers=_headers(),
        )
    assert resp.status_code == 200
    data = resp.json()
    assert data["is_binary"] is False
    assert data["is_new"] is False
    assert "wp-diff-add" in data["diff_html"]
    assert "nouveau" in data["diff_html"]


def test_preview_change_new_file(tmp_path: Path) -> None:
    project = tmp_path / "project"
    project.mkdir()
    ws_data = tmp_path / "ws_data"
    ws_data.mkdir()

    with TestClient(mainmod.app) as client:
        resp = client.post(
            "/documents/preview-change",
            json={
                "workspace_data_dir": str(ws_data),
                "project_path": str(project),
                "file_path": "nouveau.txt",
                "proposed_content": "contenu",
            },
            headers=_headers(),
        )
    assert resp.status_code == 200
    assert resp.json()["is_new"] is True


def test_preview_change_binary_docx_without_tool_args(tmp_path: Path) -> None:
    project = tmp_path / "project"
    project.mkdir()
    ws_data = tmp_path / "ws_data"
    ws_data.mkdir()
    doc = Document()
    doc.add_paragraph("Corps")
    doc.save(project / "rapport.docx")

    with TestClient(mainmod.app) as client:
        resp = client.post(
            "/documents/preview-change",
            json={
                "workspace_data_dir": str(ws_data),
                "project_path": str(project),
                "file_path": "rapport.docx",
                "proposed_content": "ignored",
            },
            headers=_headers(),
        )
    assert resp.status_code == 200
    data = resp.json()
    assert data["is_binary"] is True
    assert data["diff_html"] == ""
    assert data["message"]


def test_preview_change_docx_with_tool_args(tmp_path: Path) -> None:
    project = tmp_path / "project"
    project.mkdir()
    ws_data = tmp_path / "ws_data"
    ws_data.mkdir()
    doc = Document()
    doc.add_paragraph("Ancien paragraphe")
    doc.save(project / "rapport.docx")

    with TestClient(mainmod.app) as client:
        resp = client.post(
            "/documents/preview-change",
            json={
                "workspace_data_dir": str(ws_data),
                "project_path": str(project),
                "file_path": "rapport.docx",
                "tool_name": "write_docx",
                "tool_args": {
                    "title": "Rapport",
                    "paragraphs": ["Nouveau paragraphe"],
                },
            },
            headers=_headers(),
        )
    assert resp.status_code == 200
    data = resp.json()
    assert data["is_binary"] is False
    assert data["is_new"] is False
    assert "wp-diff-add" in data["diff_html"]
    assert "Nouveau paragraphe" in data["diff_html"]
    assert "Ancien paragraphe" in data["diff_html"]


def test_preview_change_xlsx_with_tool_args(tmp_path: Path) -> None:
    from openpyxl import Workbook

    project = tmp_path / "project"
    project.mkdir()
    ws_data = tmp_path / "ws_data"
    ws_data.mkdir()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Données"
    sheet.append(["Ancien", "Valeur"])
    workbook.save(project / "data.xlsx")
    workbook.close()

    with TestClient(mainmod.app) as client:
        resp = client.post(
            "/documents/preview-change",
            json={
                "workspace_data_dir": str(ws_data),
                "project_path": str(project),
                "file_path": "data.xlsx",
                "tool_name": "write_xlsx",
                "tool_args": {
                    "sheets": [
                        {
                            "name": "Données",
                            "rows": [["Nouveau", "Valeur"]],
                        }
                    ],
                },
            },
            headers=_headers(),
        )
    assert resp.status_code == 200
    data = resp.json()
    assert data["is_binary"] is False
    assert "wp-diff-add" in data["diff_html"]
    assert "Nouveau" in data["diff_html"]


def test_preview_change_pptx_with_tool_args(tmp_path: Path) -> None:
    project = tmp_path / "project"
    project.mkdir()
    ws_data = tmp_path / "ws_data"
    ws_data.mkdir()

    with TestClient(mainmod.app) as client:
        resp = client.post(
            "/documents/preview-change",
            json={
                "workspace_data_dir": str(ws_data),
                "project_path": str(project),
                "file_path": "pitch.pptx",
                "tool_name": "write_pptx",
                "tool_args": {
                    "theme": "improba",
                    "slides": [
                        {
                            "layout": "title",
                            "title": "Nouveau pitch",
                            "subtitle": "Vision",
                        },
                        {
                            "layout": "bullets",
                            "title": "Plan",
                            "bullets": ["Étape 1", "Étape 2"],
                        },
                    ],
                },
            },
            headers=_headers(),
        )
    assert resp.status_code == 200
    data = resp.json()
    assert data["is_binary"] is False
    assert data["is_new"] is True
    assert "wp-diff-add" in data["diff_html"]
    assert "Nouveau pitch" in data["diff_html"]
    assert "Étape 1" in data["diff_html"]
    assert "wp-slide" in data.get("preview_html", "")


def test_preview_change_pptx_applies_critique_split(tmp_path: Path) -> None:
    """L'aperçu HTML doit refléter l'auto-split (même deck que le PPTX final)."""
    project = tmp_path / "project"
    project.mkdir()
    ws_data = tmp_path / "ws_data"
    ws_data.mkdir()
    bullets = [f"Point {i} avec du texte" for i in range(12)]

    with TestClient(mainmod.app) as client:
        resp = client.post(
            "/documents/preview-change",
            json={
                "workspace_data_dir": str(ws_data),
                "project_path": str(project),
                "file_path": "dense.pptx",
                "tool_name": "write_pptx",
                "tool_args": {
                    "theme": "improba",
                    "slides": [
                        {
                            "grammar": "sequence",
                            "hierarchy": {"primary": "Dense", "secondary": bullets},
                            "density": "high",
                        }
                    ],
                },
            },
            headers=_headers(),
        )
    assert resp.status_code == 200
    html = resp.json().get("preview_html", "")
    assert html.count("wp-slide") >= 2


def test_preview_change_pptx_modify_existing(tmp_path: Path) -> None:
    from app.documents.writer import build_pptx_bytes

    project = tmp_path / "project"
    project.mkdir()
    ws_data = tmp_path / "ws_data"
    ws_data.mkdir()
    (project / "pitch.pptx").write_bytes(
        build_pptx_bytes(
            slides=[{"layout": "title", "title": "Ancien titre", "subtitle": "v1"}]
        )
    )

    with TestClient(mainmod.app) as client:
        resp = client.post(
            "/documents/preview-change",
            json={
                "workspace_data_dir": str(ws_data),
                "project_path": str(project),
                "file_path": "pitch.pptx",
                "tool_name": "write_pptx",
                "tool_args": {
                    "slides": [
                        {
                            "layout": "title",
                            "title": "Nouveau titre",
                            "subtitle": "v2",
                        }
                    ],
                },
            },
            headers=_headers(),
        )
    assert resp.status_code == 200
    data = resp.json()
    assert data["is_new"] is False
    assert data["is_binary"] is False
    assert "Ancien titre" in data["diff_html"]
    assert "Nouveau titre" in data["diff_html"]


def test_preview_change_pptx_corrupt_existing_is_graceful(tmp_path: Path) -> None:
    project = tmp_path / "project"
    project.mkdir()
    ws_data = tmp_path / "ws_data"
    ws_data.mkdir()
    (project / "broken.pptx").write_bytes(b"not-a-real-pptx")

    with TestClient(mainmod.app) as client:
        resp = client.post(
            "/documents/preview-change",
            json={
                "workspace_data_dir": str(ws_data),
                "project_path": str(project),
                "file_path": "broken.pptx",
                "tool_name": "write_pptx",
                "tool_args": {
                    "slides": [{"layout": "title", "title": "OK"}],
                },
            },
            headers=_headers(),
        )
    assert resp.status_code == 200
    data = resp.json()
    # Ancien illisible: on montre quand même le contenu proposé
    assert data["is_binary"] is False
    assert "OK" in data["diff_html"]
    assert "wp-diff-add" in data["diff_html"]


def test_preview_change_pptx_invalid_theme_is_graceful(tmp_path: Path) -> None:
    project = tmp_path / "project"
    project.mkdir()
    ws_data = tmp_path / "ws_data"
    ws_data.mkdir()

    with TestClient(mainmod.app) as client:
        resp = client.post(
            "/documents/preview-change",
            json={
                "workspace_data_dir": str(ws_data),
                "project_path": str(project),
                "file_path": "pitch.pptx",
                "tool_name": "write_pptx",
                "tool_args": {
                    "theme": "neon",
                    "slides": [{"layout": "title", "title": "X"}],
                },
            },
            headers=_headers(),
        )
    assert resp.status_code == 200
    data = resp.json()
    assert data["is_binary"] is True
    assert data["message"]


def test_preview_change_rejects_traversal(tmp_path: Path) -> None:
    project = tmp_path / "project"
    project.mkdir()
    ws_data = tmp_path / "ws_data"
    ws_data.mkdir()
    outside = tmp_path / "secret.txt"
    outside.write_text("secret", encoding="utf-8")

    with TestClient(mainmod.app) as client:
        resp = client.post(
            "/documents/preview-change",
            json={
                "workspace_data_dir": str(ws_data),
                "project_path": str(project),
                "file_path": "../secret.txt",
                "proposed_content": "hack",
            },
            headers=_headers(),
        )
    assert resp.status_code == 403
