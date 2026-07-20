"""Génération native de fichiers Office et PDF."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from app.documents.pptx_builder import (
    MAX_PPTX_SLIDES,
    PPTX_LAYOUTS,
    PPTX_THEMES,
    build_pptx_bytes,
)

__all__ = [
    "MAX_PPTX_SLIDES",
    "PPTX_LAYOUTS",
    "PPTX_THEMES",
    "build_docx_bytes",
    "build_pdf_bytes",
    "build_pptx_bytes",
    "build_xlsx_bytes",
    "require_path_extension",
]


def require_path_extension(path: str, expected_ext: str) -> None:
    """Refuse un chemin dont l'extension ne correspond pas au format réel.

    Évite notamment un .docx persisté sous un nom .pptx.
    """
    suffix = Path(path).suffix.lower()
    expected = expected_ext.lower()
    if not expected.startswith("."):
        expected = f".{expected}"
    if suffix != expected:
        raise ValueError(
            f"Le chemin doit se terminer par {expected} (reçu: {suffix or 'aucune extension'})"
        )


def build_docx_bytes(
    *,
    title: str | None = None,
    paragraphs: list[str] | None = None,
) -> bytes:
    from docx import Document

    document = Document()
    if title:
        document.add_heading(title, level=1)
    for paragraph in paragraphs or []:
        text = paragraph.strip()
        if text:
            document.add_paragraph(text)
    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


def build_xlsx_bytes(
    *,
    sheets: list[dict[str, Any]] | None = None,
) -> bytes:
    from openpyxl import Workbook

    workbook = Workbook()
    workbook.remove(workbook.active)
    sheet_defs = sheets or [{"name": "Sheet1", "rows": []}]

    for index, sheet_def in enumerate(sheet_defs):
        name = str(sheet_def.get("name") or f"Sheet{index + 1}")
        worksheet = workbook.create_sheet(title=name[:31])
        rows = sheet_def.get("rows")
        if isinstance(rows, list):
            for row in rows:
                if isinstance(row, list):
                    worksheet.append(row)

    buf = BytesIO()
    workbook.save(buf)
    return buf.getvalue()


def build_pdf_bytes(
    *,
    title: str | None = None,
    sections: list[dict[str, Any]] | None = None,
) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise RuntimeError("reportlab is not installed") from exc

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    styles = getSampleStyleSheet()
    story: list[Any] = []

    if title:
        story.append(Paragraph(title, styles["Title"]))
        story.append(Spacer(1, 12))

    for section in sections or []:
        heading = section.get("heading")
        if isinstance(heading, str) and heading.strip():
            story.append(Paragraph(heading.strip(), styles["Heading2"]))
            story.append(Spacer(1, 6))
        body = section.get("body")
        if isinstance(body, str) and body.strip():
            for line in body.strip().splitlines():
                story.append(Paragraph(line, styles["Normal"]))
            story.append(Spacer(1, 8))

    doc.build(story)
    return buf.getvalue()
